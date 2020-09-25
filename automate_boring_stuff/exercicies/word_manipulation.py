import docx

doc = docx.Document("example_word.docx")

print(len(doc.paragraphs))
# 117

print(doc.paragraphs[0].text)

# A Run object is a contiguous run of text with the same style
print(len(doc.paragraphs[1].runs))

# Getting all the text

fullText = []

for para in doc.paragraphs:
    fullText.append(para.text)

new_full_text = "\n".join(fullText)
print(new_full_text[:200])

# Writing word documents

doc = docx.Document()

doc.add_paragraph("hello world!!")
doc.save("created_word.docx")
